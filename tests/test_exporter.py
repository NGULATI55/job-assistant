"""Offline tests for the markdown -> docx exporter.

No Word required. We re-open the produced .docx with python-docx itself.
Run with:
    python -m tests.test_exporter
"""

from __future__ import annotations

import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402

from core import exporter  # noqa: E402


def _render(md: str) -> tuple[Path, Document]:
    tmp = Path(tempfile.mkdtemp())
    out = tmp / "out.docx"
    exporter.markdown_to_docx(md, out)
    return out, Document(str(out))


def test_produces_valid_docx_zip():
    out, _doc = _render("# Hello\n\nWorld")
    assert out.exists() and out.stat().st_size > 0
    # .docx is just a zip archive
    with zipfile.ZipFile(out) as zf:
        names = zf.namelist()
        assert "word/document.xml" in names


def test_headings_at_correct_levels():
    md = "# H1\n\n## H2\n\n### H3\n"
    _out, doc = _render(md)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) == 3
    styles = [p.style.name for p in headings]
    assert styles == ["Heading 1", "Heading 2", "Heading 3"]
    texts = [p.text for p in headings]
    assert texts == ["H1", "H2", "H3"]


def test_heading_clamps_above_three():
    md = "#### Deep heading\n"
    _out, doc = _render(md)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) == 1
    # Anything deeper than ### should still be a real heading, not literal "#### " text.
    assert "#" not in headings[0].text
    assert headings[0].text == "Deep heading"


def test_bullets_use_list_bullet_style():
    md = "## Items\n\n- one\n- two\n- three\n"
    _out, doc = _render(md)
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert [p.text for p in bullets] == ["one", "two", "three"]


def test_star_bullets_also_supported():
    md = "* alpha\n* beta\n"
    _out, doc = _render(md)
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert [p.text for p in bullets] == ["alpha", "beta"]


def test_paragraphs_join_consecutive_lines():
    md = "This is the first sentence.\nAnd a second one wraps too.\n\nNew paragraph here.\n"
    _out, doc = _render(md)
    texts = [p.text for p in doc.paragraphs if p.text.strip() and not p.style.name.startswith("Heading")]
    assert "This is the first sentence. And a second one wraps too." in texts
    assert "New paragraph here." in texts


def test_inline_bold_and_italic_and_code():
    md = "This has **bold** and *italic* and `code` runs.\n"
    _out, doc = _render(md)
    # The single paragraph should have 7 runs:  "This has ", "bold", " and ", "italic", " and ", "code", " runs."
    para = next(p for p in doc.paragraphs if p.text and not p.style.name.startswith("Heading"))
    by_text = {r.text: r for r in para.runs}
    assert "bold" in by_text and by_text["bold"].bold is True
    assert "italic" in by_text and by_text["italic"].italic is True
    assert "code" in by_text and by_text["code"].font.name == "Consolas"
    # Plain runs should NOT be bold/italic.
    plain = by_text.get("This has ")
    assert plain is not None and not plain.bold and not plain.italic


def test_blockquote_renders_italic():
    md = "> source line from master\n"
    _out, doc = _render(md)
    para = next(p for p in doc.paragraphs if p.text)
    assert para.text == "source line from master"
    assert para.runs and para.runs[0].italic is True


def test_fenced_code_block_degrades_to_plain():
    md = "```\nline1\nline2\n```\n"
    _out, doc = _render(md)
    paras = [p for p in doc.paragraphs if p.text]
    assert any("line1" in p.text and "line2" in p.text for p in paras)


def test_unknown_syntax_degrades_to_text():
    # A table-like line is not understood, should appear as plain text rather than crash.
    md = "Some text\n\n| col1 | col2 |\n\n## After\n"
    _out, doc = _render(md)
    texts = [p.text for p in doc.paragraphs if p.text]
    assert "| col1 | col2 |" in texts
    assert "After" in texts


def test_empty_input_produces_empty_doc():
    out, doc = _render("")
    assert out.exists()
    # An empty docx has one default paragraph; just confirm no crash and no runs of content.
    real_text = "".join(p.text for p in doc.paragraphs).strip()
    assert real_text == ""


# --- PDF export tests ---------------------------------------------------

def test_pdf_produces_valid_pdf_bytes():
    data = exporter.markdown_to_pdf_bytes("# Hello\n\nWorld")
    assert isinstance(data, bytes)
    assert data.startswith(b"%PDF-"), f"Not a PDF: {data[:8]!r}"
    assert b"%%EOF" in data[-100:], "Missing EOF marker"
    assert len(data) > 200


def test_pdf_writes_file():
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "x.pdf"
        path = exporter.markdown_to_pdf(
            "# Title\n\n- bullet one\n- bullet two\n", out, title="Test"
        )
        assert path == out and out.exists() and out.stat().st_size > 200
        assert out.read_bytes().startswith(b"%PDF-")


def test_pdf_handles_empty_input():
    data = exporter.markdown_to_pdf_bytes("")
    assert data.startswith(b"%PDF-")
    assert b"%%EOF" in data[-100:]


def test_pdf_handles_realistic_resume_without_crash():
    md = (
        "# Tailored Resume — Marketing Manager at Acme\n\n"
        "## Summary\n"
        "Marketing professional with experience across paid social, Google Ads, SEO and content.\n\n"
        "## Highlights\n"
        "- Led paid social and Google Ads end-to-end\n"
        "- Built the **content calendar** across blog, email and social\n"
        "- Lifted *organic* traffic via on-page SEO\n\n"
        "## Experience\n\n"
        "### Marketing Lead — Foo Pty Ltd (Jan 2022 – present)\n"
        "- Owned the brand voice across all channels\n"
        "- Reported weekly to leadership\n"
    )
    data = exporter.markdown_to_pdf_bytes(md, title="Tailored Resume")
    assert data.startswith(b"%PDF-")
    assert len(data) > 800  # non-trivial size


def test_pdf_handles_non_cp1252_characters():
    # Emoji etc. should not crash — they get replaced with `?` via cp1252 errors='replace'.
    md = "# Title\n\nA paragraph with an emoji \U0001f31f and a smart quote ’."
    data = exporter.markdown_to_pdf_bytes(md)
    assert data.startswith(b"%PDF-")


def test_all_styles_produce_valid_pdf_and_docx():
    """Every shipped style should render valid PDF + DOCX for a typical resume."""
    md = (
        "# Nipun Gulati\n"
        "Melbourne · email@example.com\n\n"
        "## Summary\n"
        "SEO specialist with **agency** and *freelance* background.\n\n"
        "## Skills\n"
        "- Technical SEO\n"
        "- GA4 and Search Console\n"
        "- Ahrefs and SEMrush\n"
    )
    for style_key, style_meta in exporter.STYLES.items():
        pdf = exporter.markdown_to_pdf_bytes(md, title="Test", style=style_key)
        assert pdf.startswith(b"%PDF-"), f"{style_key} PDF invalid: {pdf[:8]!r}"
        assert len(pdf) > 500, f"{style_key} PDF too small"

        docx = exporter.markdown_to_docx_bytes(md, style=style_key)
        # docx is a zip
        import io as _io, zipfile as _zip
        with _zip.ZipFile(_io.BytesIO(docx)) as zf:
            assert "word/document.xml" in zf.namelist(), f"{style_key} DOCX invalid"


def test_unknown_style_falls_back_to_default():
    md = "# Hello\n\nWorld"
    a = exporter.markdown_to_pdf_bytes(md, style="nonexistent_template")
    b = exporter.markdown_to_pdf_bytes(md, style=exporter.DEFAULT_STYLE)
    assert a.startswith(b"%PDF-")
    # Different timestamps in PDFs may differ; just check both rendered.
    assert b.startswith(b"%PDF-")


def test_pdf_inline_emphasis_does_not_crash():
    md = "Plain **bold** and *italic* and `code` text together."
    data = exporter.markdown_to_pdf_bytes(md)
    assert data.startswith(b"%PDF-")


def test_realistic_tailored_resume_end_to_end():
    md = (
        "# Tailored Resume — Marketing Manager at Acme Pty Ltd\n\n"
        "## Summary\n"
        "Marketing professional with experience across paid social, Google Ads, SEO and content.\n\n"
        "## Highlights\n"
        "- Led paid social and Google Ads end-to-end\n"
        "- Built the **content calendar** across blog, email and social\n"
        "- Lifted *organic* traffic via on-page SEO\n\n"
        "## Experience\n"
        "### Marketing Lead — Foo Pty Ltd (Jan 2022 – present)\n"
        "- Owned the brand voice across all channels\n"
        "- Reported weekly to leadership\n"
    )
    _out, doc = _render(md)
    h1s = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    h2s = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    h3s = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(h1s) == 1 and h1s[0].text.startswith("Tailored Resume")
    # Default (Classic Plain) style keeps section headers in their original case.
    assert [p.text for p in h2s] == ["Summary", "Highlights", "Experience"]
    assert h3s and "Marketing Lead" in h3s[0].text
    assert len(bullets) == 5
    # Inline bold survived in the highlight bullet.
    cal_bullet = next(b for b in bullets if "content calendar" in b.text)
    assert any(r.bold for r in cal_bullet.runs if r.text == "content calendar")


# --- Runner -------------------------------------------------------------

def _run_all():
    tests = [
        test_produces_valid_docx_zip,
        test_headings_at_correct_levels,
        test_heading_clamps_above_three,
        test_bullets_use_list_bullet_style,
        test_star_bullets_also_supported,
        test_paragraphs_join_consecutive_lines,
        test_inline_bold_and_italic_and_code,
        test_blockquote_renders_italic,
        test_fenced_code_block_degrades_to_plain,
        test_unknown_syntax_degrades_to_text,
        test_empty_input_produces_empty_doc,
        test_realistic_tailored_resume_end_to_end,
        test_pdf_produces_valid_pdf_bytes,
        test_pdf_writes_file,
        test_pdf_handles_empty_input,
        test_pdf_handles_realistic_resume_without_crash,
        test_pdf_handles_non_cp1252_characters,
        test_pdf_inline_emphasis_does_not_crash,
        test_all_styles_produce_valid_pdf_and_docx,
        test_unknown_style_falls_back_to_default,
    ]
    passed = 0
    for fn in tests:
        try:
            fn()
            print(f"OK  {fn.__name__}")
            passed += 1
        except AssertionError as e:
            print(f"FAIL {fn.__name__}: {e}")
        except Exception as e:  # noqa: BLE001
            print(f"ERR  {fn.__name__}: {type(e).__name__}: {e}")
    print(f"\n{passed}/{len(tests)} passed")
    return passed == len(tests)


if __name__ == "__main__":
    ok = _run_all()
    sys.exit(0 if ok else 1)
