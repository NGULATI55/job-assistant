"""Offline tests for the resume library.

Focus: filename sanitisation (no path traversal), list ordering, round-trip.
Run with:
    python -m tests.test_resume_loader
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from core import resume_loader  # noqa: E402


def test_list_resumes_returns_sorted_md_only():
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        (root / "b.md").write_text("b")
        (root / "a.md").write_text("a")
        (root / "ignore.txt").write_text("nope")
        (root / "sub").mkdir()
        (root / "sub" / "c.md").write_text("nested — should NOT be listed")
        names = resume_loader.list_resumes(root)
        assert names == ["a.md", "b.md"]


def test_list_resumes_empty_dir():
    with tempfile.TemporaryDirectory() as td:
        assert resume_loader.list_resumes(Path(td)) == []


def test_list_resumes_missing_dir():
    assert resume_loader.list_resumes(Path("nonexistent_folder_xyz")) == []


def test_safe_resume_name_basic():
    assert resume_loader.safe_resume_name("nipun.md") == "nipun.md"
    assert resume_loader.safe_resume_name("Nipun Gulati.md") == "Nipun Gulati.md"


def test_safe_resume_name_strips_path():
    assert resume_loader.safe_resume_name("../../etc/passwd.md") == "passwd.md"
    assert resume_loader.safe_resume_name("/abs/path/foo.md") == "foo.md"
    assert resume_loader.safe_resume_name("C:\\Users\\bunny\\evil.md") == "evil.md"


def test_safe_resume_name_appends_md():
    assert resume_loader.safe_resume_name("nipun") == "nipun.md"
    # Any input extension is normalized to lowercase .md (so "john.MD", "john.pdf", "john.docx"
    # all become "john.md").
    assert resume_loader.safe_resume_name("john.MD") == "john.md"


def test_safe_resume_name_rejects_garbage():
    for bad in ("", "   ", ".", "..", "/", "\\"):
        try:
            resume_loader.safe_resume_name(bad)
        except ValueError:
            continue
        raise AssertionError(f"Expected ValueError for {bad!r}")


def test_save_uploaded_resume_round_trip():
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        path = resume_loader.save_uploaded_resume(root, "alice.md", b"# Alice\n\nHi")
        assert path == (root / "alice.md").resolve()
        assert path.read_text(encoding="utf-8") == "# Alice\n\nHi"


def test_save_uploaded_resume_creates_dir():
    with tempfile.TemporaryDirectory() as td:
        root = Path(td) / "nested" / "resumes"
        assert not root.exists()
        path = resume_loader.save_uploaded_resume(root, "x.md", b"x")
        assert path.exists()
        assert root.exists()


def test_save_uploaded_resume_blocks_traversal():
    with tempfile.TemporaryDirectory() as td:
        root = Path(td) / "resumes"
        root.mkdir()
        # Even with a traversal-ish name, the basename is extracted and the file
        # ends up safely inside `root`.
        path = resume_loader.save_uploaded_resume(root, "../evil.md", b"data")
        assert path.parent.resolve() == root.resolve()
        assert path.name == "evil.md"
        # The parent dir should NOT contain `evil.md` (i.e. the traversal didn't escape).
        assert not (root.parent / "evil.md").exists()


def test_save_uploaded_resume_rejects_empty_name():
    with tempfile.TemporaryDirectory() as td:
        try:
            resume_loader.save_uploaded_resume(Path(td), "", b"x")
        except ValueError:
            return
        raise AssertionError("Expected ValueError for empty filename")


def test_load_resume_text_missing_returns_empty():
    assert resume_loader.load_resume_text(Path("nonexistent.md")) == ""


def test_load_resume_text_reads_utf8():
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "x.md"
        p.write_text("héllo — world", encoding="utf-8")
        assert resume_loader.load_resume_text(p) == "héllo — world"


# --- New: multi-format extraction --------------------------------------

def test_extract_text_handles_md_passthrough():
    text = resume_loader.extract_text("foo.md", b"# Hello\n\nWorld")
    assert text == "# Hello\n\nWorld"


def test_extract_text_handles_txt():
    text = resume_loader.extract_text("foo.txt", b"plain text content\nline 2")
    assert text == "plain text content\nline 2"


def test_extract_text_rejects_unsupported_extension():
    try:
        resume_loader.extract_text("foo.rtf", b"data")
    except ValueError as e:
        assert "Unsupported" in str(e)
        return
    raise AssertionError("Expected ValueError for .rtf")


def test_extract_text_handles_pdf():
    """Build a minimal PDF in memory and verify text extraction."""
    try:
        import fitz
    except ImportError:
        return  # skip if pymupdf missing
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "John Doe\nSenior Engineer\nSydney")
    pdf_bytes = doc.tobytes()
    doc.close()
    text = resume_loader.extract_text("john.pdf", pdf_bytes)
    assert "John Doe" in text
    assert "Senior Engineer" in text
    assert "Sydney" in text


def test_extract_text_handles_docx():
    """Build a minimal DOCX in memory and verify headings + bullets are preserved."""
    from docx import Document
    import io as _io
    d = Document()
    d.add_heading("Jane Smith", level=1)
    d.add_heading("Summary", level=2)
    d.add_paragraph("Marketing professional with 5 years experience.")
    d.add_heading("Skills", level=2)
    d.add_paragraph("Google Ads", style="List Bullet")
    d.add_paragraph("SEO", style="List Bullet")
    buf = _io.BytesIO()
    d.save(buf)
    text = resume_loader.extract_text("jane.docx", buf.getvalue())
    assert "# Jane Smith" in text
    assert "## Summary" in text
    assert "## Skills" in text
    assert "- Google Ads" in text
    assert "- SEO" in text


def test_save_uploaded_resume_extracts_pdf_to_md():
    """Uploaded PDF should be stored as <name>.md with extracted text."""
    try:
        import fitz
    except ImportError:
        return
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Alice Carter\nProduct Manager")
    pdf_bytes = doc.tobytes()
    doc.close()
    with tempfile.TemporaryDirectory() as td:
        path = resume_loader.save_uploaded_resume(Path(td), "alice.pdf", pdf_bytes)
        assert path.suffix == ".md"
        assert path.name == "alice.md"
        content = path.read_text(encoding="utf-8")
        assert "Alice Carter" in content
        assert "Product Manager" in content


def test_save_uploaded_resume_rejects_unsupported_format():
    with tempfile.TemporaryDirectory() as td:
        try:
            resume_loader.save_uploaded_resume(Path(td), "foo.rtf", b"data")
        except ValueError as e:
            assert "Unsupported" in str(e)
            return
    raise AssertionError("Expected ValueError for .rtf")


def test_safe_resume_name_strips_any_original_extension():
    assert resume_loader.safe_resume_name("john.pdf") == "john.md"
    assert resume_loader.safe_resume_name("jane.docx") == "jane.md"
    assert resume_loader.safe_resume_name("notes.txt") == "notes.md"
    assert resume_loader.safe_resume_name("existing.md") == "existing.md"


# --- Runner -------------------------------------------------------------

def _run_all():
    tests = [
        test_list_resumes_returns_sorted_md_only,
        test_list_resumes_empty_dir,
        test_list_resumes_missing_dir,
        test_safe_resume_name_basic,
        test_safe_resume_name_strips_path,
        test_safe_resume_name_appends_md,
        test_safe_resume_name_rejects_garbage,
        test_save_uploaded_resume_round_trip,
        test_save_uploaded_resume_creates_dir,
        test_save_uploaded_resume_blocks_traversal,
        test_save_uploaded_resume_rejects_empty_name,
        test_load_resume_text_missing_returns_empty,
        test_load_resume_text_reads_utf8,
        test_extract_text_handles_md_passthrough,
        test_extract_text_handles_txt,
        test_extract_text_rejects_unsupported_extension,
        test_extract_text_handles_pdf,
        test_extract_text_handles_docx,
        test_save_uploaded_resume_extracts_pdf_to_md,
        test_save_uploaded_resume_rejects_unsupported_format,
        test_safe_resume_name_strips_any_original_extension,
    ]
    passed = 0
    for fn in tests:
        try:
            fn()
            print(f"OK  {fn.__name__}")
            passed += 1
        except AssertionError as e:
            print(f"FAIL {fn.__name__}: {e}")
        except Exception as e:  # noqa: BLE001
            print(f"ERR  {fn.__name__}: {type(e).__name__}: {e}")
    print(f"\n{passed}/{len(tests)} passed")
    return passed == len(tests)


if __name__ == "__main__":
    ok = _run_all()
    sys.exit(0 if ok else 1)
