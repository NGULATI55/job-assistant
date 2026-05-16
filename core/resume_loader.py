"""Resume library at data/resumes/. Accepts uploads in PDF, DOCX, TXT, or MD.

Anything uploaded is extracted to plain text and stored as `<name>.md` so the
rest of the pipeline (the LLM tailoring prompt) only has to deal with text.
"""

from __future__ import annotations

import io
from pathlib import Path


_BUILTIN_EXAMPLE_PATH = Path(__file__).resolve().parents[1] / "data" / "resumes" / "example.md"

_BUILTIN_EXAMPLE_FALLBACK = (
    "# Example Resume — Replace this with your details\n\n"
    "> This is a starter template. Either edit this content here, or upload your own resume.\n\n"
    "## Contact\n- Name:\n- Location:\n- Email:\n- Phone:\n\n"
    "## Summary\nA short 2 to 3 sentence summary in your own voice.\n\n"
    "## Experience\n\n### Role title — Company (Month YYYY – Month YYYY)\n- Achievement with a number.\n- Achievement with a number.\n\n"
    "## Skills\n- Skill 1\n- Skill 2\n\n## Education\n- Degree, Institution, Year\n"
)


SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


def builtin_example() -> str:
    if _BUILTIN_EXAMPLE_PATH.exists():
        try:
            return _BUILTIN_EXAMPLE_PATH.read_text(encoding="utf-8")
        except OSError:
            pass
    return _BUILTIN_EXAMPLE_FALLBACK


def list_resumes(resumes_dir: Path) -> list[str]:
    if not resumes_dir.exists():
        return []
    return sorted(p.name for p in resumes_dir.glob("*.md") if p.is_file())


def load_resume_text(resume_path: Path) -> str:
    if not resume_path.exists():
        return ""
    return resume_path.read_text(encoding="utf-8")


def safe_resume_name(raw_name: str) -> str:
    """Strip path, drop original extension, force .md.

    "Jane CV.pdf" -> "Jane CV.md"
    "../../etc/passwd.pdf" -> "passwd.md"
    """
    base = Path(raw_name).name.strip()
    if not base or base in (".", ".."):
        raise ValueError("Invalid filename.")
    stem = Path(base).stem.strip() or "resume"
    return f"{stem}.md"


# --- Text extraction ---------------------------------------------------

def extract_text(raw_name: str, data: bytes) -> str:
    """Extract plain text from a resume in any supported format."""
    ext = Path(raw_name).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    if ext in (".txt", ".md"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(
        f"Unsupported resume format: '{ext}'. "
        f"Use one of: {', '.join(SUPPORTED_EXTENSIONS)}"
    )


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ValueError(
            "PDF extraction requires PyMuPDF. Install with: pip install pymupdf"
        ) from e
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        pages = [page.get_text("text") for page in doc]
    finally:
        doc.close()
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError(
            "Could not extract any text from this PDF. "
            "If it's an image-only / scanned PDF, save it as text first."
        )
    return text


def _extract_docx(data: bytes) -> str:
    """Extract text from a DOCX file, preserving headings + bullets as markdown."""
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError(
            "DOCX extraction requires python-docx. Install with: pip install python-docx"
        ) from e
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if not text.strip():
            parts.append("")
            continue
        style = para.style.name if para.style else "Normal"
        if "Heading 1" in style or style == "Title":
            parts.append(f"# {text}")
        elif "Heading 2" in style:
            parts.append(f"## {text}")
        elif "Heading 3" in style or "Heading 4" in style:
            parts.append(f"### {text}")
        elif style == "List Bullet" or style.startswith("List "):
            parts.append(f"- {text}")
        elif text.startswith(("•", "·", "●", "▪")):
            parts.append("- " + text.lstrip("•·●▪").strip())
        else:
            parts.append(text)
    # Also walk tables (resumes sometimes use tables for contact / 2-col layouts).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t and t not in parts:
                        parts.append(t)
    out = "\n".join(parts).strip()
    if not out:
        raise ValueError("Could not extract any text from this DOCX.")
    return out


# --- Save (with extraction on upload) ----------------------------------

def save_uploaded_resume(resumes_dir: Path, raw_name: str, data: bytes) -> Path:
    """Save uploaded resume in any supported format. Always stored as .md text.

    Defends against path traversal: basename only, then validated against the root.
    """
    ext = Path(raw_name).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported format '{ext}'. "
            f"Use one of: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    text = extract_text(raw_name, data)
    name = safe_resume_name(raw_name)
    resumes_dir.mkdir(parents=True, exist_ok=True)
    target = (resumes_dir / name).resolve()
    root = resumes_dir.resolve()
    if root not in target.parents:
        raise ValueError("Refusing to write outside the resumes directory.")
    target.write_text(text, encoding="utf-8")
    return target
