"""Markdown -> .docx converter for the tailored resume.

Deliberately simple and robust. Supports:
- Headings #, ##, ### (mapped to Word Heading 1/2/3)
- Plain paragraphs (consecutive non-blank lines join with a space)
- Bullet points starting with "- " or "* "
- Blockquote lines starting with "> " (rendered as italic paragraph)
- Inline **bold**, *italic*, and `code`

Anything we don't recognise degrades to plain text. Output is ATS-friendly:
Calibri defaults, no tables, no images, no fancy styles.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from fpdf import FPDF


# ============================================================================
# Style templates. All single-column, ATS-readable. They differ in colour,
# typography and spacing, not structure.
# ============================================================================

STYLES: dict[str, dict] = {
    "editorial_bold": {
        "name": "Editorial Bold",
        "description": "Huge 36pt name with warm-orange accent on the first word, thick 1mm rule. Single column, magazine-cover energy. ATS-safe.",
        "pdf_body_font": "Helvetica", "pdf_heading_font": "Helvetica",
        "pdf_h1_size": 34, "pdf_h2_size": 13, "pdf_h3_size": 11.5, "pdf_body_size": 10.5, "pdf_line_h": 5.4,
        "docx_body_font": "Calibri", "docx_heading_font": "Calibri",
        "docx_h1_size_pt": 34, "docx_h2_size_pt": 13, "docx_h3_size_pt": 11.5, "docx_body_size_pt": 10.5,
        "color_h1": (20, 20, 20), "color_h2": (20, 20, 20), "color_h3": (90, 90, 90), "color_body": (40, 40, 40),
        "h1_case": "normal", "h2_case": "normal", "h2_letter_spacing_pt": 0,
        "rule_under_h1": True, "rule_under_h2": True,
        "rule_color": (20, 20, 20), "rule_thickness": 1.0,
        "h2_rule_color": (225, 225, 225),
        "bullet": "•",
        "h1_accent_first_word": True,
        "h1_accent_color": (232, 152, 39),  # warm orange
        "h1_align": "left",
        "banner": None, "layout": "single",
        "page_fill_color": None, "top_rule": None,
    },
    "executive_banner": {
        "name": "Executive Banner",
        "description": "Full-width navy banner with centred gold-serif name. Two-column body, Times serif, gold accent rules. Premium consultancy feel.",
        "pdf_body_font": "Times", "pdf_heading_font": "Times",
        "pdf_h1_size": 32, "pdf_h2_size": 11.5, "pdf_h3_size": 12, "pdf_body_size": 10.5, "pdf_line_h": 5.2,
        "docx_body_font": "Cambria", "docx_heading_font": "Cambria",
        "docx_h1_size_pt": 32, "docx_h2_size_pt": 11.5, "docx_h3_size_pt": 12, "docx_body_size_pt": 10.5,
        "color_h1": (201, 169, 110),  # gold inside banner
        "color_h2": (26, 58, 108),     # navy section headers
        "color_h3": (20, 20, 20), "color_body": (25, 25, 25),
        "h1_case": "normal",
        "h2_case": "upper", "h2_letter_spacing_pt": 2.0,
        "rule_under_h1": False, "rule_under_h2": True,
        "rule_color": (201, 169, 110), "h2_rule_color": (201, 169, 110),
        "bullet": "•",
        "h1_accent_first_word": False,
        "h1_align": "center",
        "banner": {"fill_color": (26, 58, 108), "height_mm": 46, "padding_top_mm": 16},
        "layout": "two_column",
        "sidebar_sections": {"skills", "education", "tools", "languages", "contact",
                              "certifications", "awards", "other", "interests", "references"},
        "sidebar_width_ratio": 0.32,
        "page_fill_color": None, "top_rule": None,
    },
    "grid_modern": {
        "name": "Grid Modern",
        "description": "Cream page with thick purple rules above and below a centred bold name. Two-column body, all-caps section headers. Designer feel without clutter.",
        "pdf_body_font": "Helvetica", "pdf_heading_font": "Helvetica",
        "pdf_h1_size": 30, "pdf_h2_size": 12, "pdf_h3_size": 11.5, "pdf_body_size": 10, "pdf_line_h": 4.9,
        "docx_body_font": "Calibri", "docx_heading_font": "Calibri",
        "docx_h1_size_pt": 30, "docx_h2_size_pt": 12, "docx_h3_size_pt": 11.5, "docx_body_size_pt": 10,
        "color_h1": (15, 15, 15), "color_h2": (15, 15, 15), "color_h3": (15, 15, 15), "color_body": (35, 35, 35),
        "h1_case": "normal",
        "h2_case": "upper", "h2_letter_spacing_pt": 1.6,
        "rule_under_h1": True, "rule_under_h2": False,
        "rule_color": (108, 70, 184), "rule_thickness": 2.0,  # purple
        "h2_rule_color": (220, 220, 220),
        "bullet": "•",
        "h1_accent_first_word": False,
        "h1_align": "center",
        "banner": None,
        "layout": "two_column",
        "sidebar_sections": {"skills", "education", "tools", "languages", "contact",
                              "certifications", "awards", "other", "interests", "references"},
        "sidebar_width_ratio": 0.45,
        "page_fill_color": (251, 249, 245),  # cream
        "top_rule": {"color": (108, 70, 184), "thickness_mm": 2.0, "y_offset_mm": 12},
    },
    "minimal_sidebar": {
        "name": "Minimal Sidebar",
        "description": "Clean two-column, bright blue uppercase section headers with matching thin rules. Disciplined whitespace, no flourish.",
        "pdf_body_font": "Helvetica", "pdf_heading_font": "Helvetica",
        "pdf_h1_size": 26, "pdf_h2_size": 11.5, "pdf_h3_size": 11, "pdf_body_size": 10, "pdf_line_h": 4.9,
        "docx_body_font": "Calibri", "docx_heading_font": "Calibri",
        "docx_h1_size_pt": 26, "docx_h2_size_pt": 11.5, "docx_h3_size_pt": 11, "docx_body_size_pt": 10,
        "color_h1": (20, 20, 20),
        "color_h2": (26, 115, 232),    # bright blue
        "color_h3": (20, 20, 20), "color_body": (35, 35, 35),
        "h1_case": "normal", "h2_case": "upper", "h2_letter_spacing_pt": 1.4,
        "rule_under_h1": False, "rule_under_h2": True,
        "rule_color": (26, 115, 232), "h2_rule_color": (26, 115, 232),
        "bullet": "•",
        "h1_accent_first_word": False,
        "h1_align": "left",
        "banner": None,
        "layout": "two_column",
        "sidebar_sections": {"skills", "education", "tools", "languages", "contact",
                              "certifications", "awards", "other", "interests", "references"},
        "sidebar_width_ratio": 0.35,
        "page_fill_color": None, "top_rule": None,
    },
}

DEFAULT_STYLE = "editorial_bold"


def _get_style(style: str | None) -> dict:
    return STYLES.get(style or DEFAULT_STYLE, STYLES[DEFAULT_STYLE])


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.+)$")
_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
_FENCE_RE = re.compile(r"^```")

# Inline tokens: **bold**, *italic*, `code`. Order matters — bold before italic.
_INLINE_RE = re.compile(r"(\*\*[^*\n]+\*\*|\*[^*\n]+\*|`[^`\n]+`)")


def markdown_to_docx_bytes(md: str, style: str | None = None) -> bytes:
    """Render Markdown to .docx bytes in memory. Useful for download-only flows."""
    doc = _build_doc(md, _get_style(style))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_docx(md: str, out_path: Path, style: str | None = None) -> Path:
    """Render Markdown to a .docx file at out_path. Returns out_path."""
    doc = _build_doc(md, _get_style(style))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _set_paragraph_bottom_border(paragraph, color_rgb: tuple[int, int, int], size: int = 6) -> None:
    """Add a thin horizontal rule under a docx paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*color_rgb))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_run_letter_spacing(run, spacing_pt: float) -> None:
    """Set character spacing on a run (Word's <w:spacing w:val=N/> in twentieths of a point)."""
    rPr = run._element.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    # Word measures spacing in 1/20 of a point.
    spc.set(qn("w:val"), str(int(spacing_pt * 20)))
    rPr.append(spc)


def _apply_run_style(run, font_name: str, size_pt: float, color_rgb: tuple[int, int, int]) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)


def _add_docx_banner(doc, name_text: str, s: dict) -> None:
    """Add a one-cell shaded table at the top simulating a colour banner header."""
    banner = s["banner"]
    fill = banner["fill_color"]
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Cell shading via direct XML
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*fill))
    tcPr.append(shd)
    # Padding inside cell
    tcMar = OxmlElement("w:tcMar")
    for side, twips in (("top", 360), ("bottom", 360), ("left", 200), ("right", 200)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(twips))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)
    # H1 paragraph inside cell
    p = cell.paragraphs[0]
    case = s.get("h1_case", "normal")
    display = name_text.upper() if case == "upper" else name_text
    _add_inline_runs(p, display)
    for run in p.runs:
        _apply_run_style(run, s["docx_heading_font"], s["docx_h1_size_pt"], s["color_h1"])
    if s.get("h1_align") == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Spacer after banner
    doc.add_paragraph()


def _add_docx_heading(doc_or_cell, level: int, raw_text: str, s: dict) -> None:
    """Add a heading with style-specific colour, font, case, letter spacing, rule."""
    case_key = "h1_case" if level == 1 else "h2_case" if level == 2 else "h3_case"
    display = raw_text.upper() if s.get(case_key, "normal") == "upper" else raw_text

    p = doc_or_cell.add_paragraph()
    p.style = doc_or_cell.styles["Heading 1"] if hasattr(doc_or_cell, "styles") else p.style
    # Note: when adding to a table cell we don't have direct .styles access; default fine.
    # Apply heading style by level - we use add_paragraph + style name string.
    try:
        p.style = (doc_or_cell.styles["Heading 1"] if level == 1 else
                   doc_or_cell.styles["Heading 2"] if level == 2 else
                   doc_or_cell.styles["Heading 3"])
    except (KeyError, AttributeError):
        # In a cell we may need a different approach; fall back to no style.
        pass

    size_pt = {1: s["docx_h1_size_pt"], 2: s["docx_h2_size_pt"], 3: s["docx_h3_size_pt"]}[level]
    color = {1: s["color_h1"], 2: s["color_h2"], 3: s["color_h3"]}[level]
    spacing_key = "h1_letter_spacing_pt" if level == 1 else "h2_letter_spacing_pt" if level == 2 else None
    spacing_pt = s.get(spacing_key, 0) if spacing_key else 0

    if level == 1 and s.get("h1_accent_first_word") and " " in display and s.get("h1_align", "left") != "center":
        first_word, rest = display.split(" ", 1)
        accent = s.get("h1_accent_color", (232, 152, 39))
        r1 = p.add_run(first_word + " ")
        _apply_run_style(r1, s["docx_heading_font"], size_pt, accent)
        r1.bold = True
        r2 = p.add_run(rest)
        _apply_run_style(r2, s["docx_heading_font"], size_pt, color)
        r2.bold = True
        if spacing_pt:
            _set_run_letter_spacing(r1, spacing_pt)
            _set_run_letter_spacing(r2, spacing_pt)
    else:
        _add_inline_runs(p, display)
        for run in p.runs:
            _apply_run_style(run, s["docx_heading_font"], size_pt, color)
            if spacing_pt:
                _set_run_letter_spacing(run, spacing_pt)

    if level == 1 and s.get("h1_align") == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if level == 1 and s.get("rule_under_h1"):
        _set_paragraph_bottom_border(p, s.get("rule_color", (200, 200, 200)), size=8)
    elif level == 2 and s.get("rule_under_h2"):
        _set_paragraph_bottom_border(p, s.get("h2_rule_color", (200, 200, 200)), size=4)


def _add_docx_paragraph(doc_or_cell, text: str, s: dict) -> None:
    p = doc_or_cell.add_paragraph()
    _add_inline_runs(p, text)
    for run in p.runs:
        if run.font.name != "Consolas":
            run.font.name = s["docx_body_font"]
        run.font.size = Pt(s["docx_body_size_pt"])
        run.font.color.rgb = RGBColor(*s["color_body"])


def _add_docx_bullet(doc_or_cell, text: str, s: dict) -> None:
    try:
        p = doc_or_cell.add_paragraph(style="List Bullet")
    except (KeyError, AttributeError):
        # In a cell, "List Bullet" may not resolve; fall back to manual bullet glyph.
        p = doc_or_cell.add_paragraph()
        p.add_run(f"{s.get('bullet', '•')}  ")
    _add_inline_runs(p, text)
    for run in p.runs:
        if run.font.name != "Consolas":
            run.font.name = s["docx_body_font"]
        run.font.size = Pt(s["docx_body_size_pt"])
        run.font.color.rgb = RGBColor(*s["color_body"])


def _add_docx_quote(doc_or_cell, text: str, s: dict) -> None:
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.name = s["docx_body_font"]
    run.font.size = Pt(s["docx_body_size_pt"])
    run.font.color.rgb = RGBColor(*s["color_body"])


def _render_block_docx(container, blk: tuple, s: dict) -> None:
    """Render a parsed block into a doc or a table cell."""
    kind = blk[0]
    if kind == "h":
        _add_docx_heading(container, blk[1], blk[2], s)
    elif kind == "p":
        _add_docx_paragraph(container, blk[1], s)
    elif kind == "bullets":
        for item in blk[1]:
            _add_docx_bullet(container, item, s)
    elif kind == "quote":
        _add_docx_quote(container, blk[1], s)
    elif kind == "code":
        p = container.add_paragraph()
        run = p.add_run(blk[1])
        run.font.name = "Consolas"


def _build_doc_two_column(doc, md: str, s: dict) -> None:
    blocks = _parse_md_blocks(md)
    sidebar_names = s.get("sidebar_sections", set())

    # H1 spans full width before columns
    if blocks and blocks[0][0] == "h" and blocks[0][1] == 1:
        _add_docx_heading(doc, blocks[0][1], blocks[0][2], s)
        blocks = blocks[1:]

    main_blocks, sidebar_blocks = _route_blocks_to_columns(blocks, sidebar_names)

    # 1-row, 2-column table for the body
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    sidebar_w_ratio = s.get("sidebar_width_ratio", 0.34)
    # A4 usable width is ~17cm. Allocate accordingly.
    from docx.shared import Cm
    total_cm = 17
    sidebar_cm = total_cm * sidebar_w_ratio
    main_cm = total_cm - sidebar_cm
    table.columns[0].width = Cm(main_cm)
    table.columns[1].width = Cm(sidebar_cm)
    for cell, width in ((table.cell(0, 0), Cm(main_cm)), (table.cell(0, 1), Cm(sidebar_cm))):
        cell.width = width

    main_cell = table.cell(0, 0)
    sidebar_cell = table.cell(0, 1)
    # Clear empty default paragraphs in cells
    for cell in (main_cell, sidebar_cell):
        # The cell starts with one empty paragraph; remove it
        if cell.paragraphs and not cell.paragraphs[0].text:
            cell._element.remove(cell.paragraphs[0]._element)

    for blk in main_blocks:
        _render_block_docx(main_cell, blk, s)
    for blk in sidebar_blocks:
        _render_block_docx(sidebar_cell, blk, s)


def _build_doc(md: str, s: dict):
    doc = Document()
    # Set the document default font on the "Normal" style so all body text inherits it.
    normal = doc.styles["Normal"]
    normal.font.name = s["docx_body_font"]
    normal.font.size = Pt(s["docx_body_size_pt"])
    normal.font.color.rgb = RGBColor(*s["color_body"])

    # Banner header (Tracy-Hall-style) takes the H1 out of the main flow.
    banner = s.get("banner")
    if banner:
        # Find H1 text in markdown to render inside banner.
        h1_text = ""
        for line in (md or "").splitlines():
            m = _HEADING_RE.match(line.rstrip())
            if m and len(m.group(1)) == 1:
                h1_text = m.group(2).strip()
                break
        if h1_text:
            _add_docx_banner(doc, h1_text, s)
            # Strip the H1 from the markdown so it doesn't render again
            md = md.replace(f"# {h1_text}", "", 1)

    # Two-column dispatcher
    if s.get("layout") == "two_column":
        _build_doc_two_column(doc, md, s)
        return doc

    lines = (md or "").splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip()

        # Skip blank lines.
        if not line.strip():
            i += 1
            continue

        # Fenced code block — pass-through as plain paragraphs, no formatting.
        if _FENCE_RE.match(line):
            i += 1
            code_lines: list[str] = []
            while i < n and not _FENCE_RE.match(lines[i].rstrip()):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
            if i < n:
                i += 1  # consume closing fence
            continue

        # Heading.
        m = _HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 3)  # clamp to Heading 1..3
            text = m.group(2).strip()

            # Case transform
            case_key = "h1_case" if level == 1 else "h2_case" if level == 2 else "h3_case"
            if s.get(case_key, "normal") == "upper":
                text = text.upper()

            p = doc.add_heading(level=level)
            _add_inline_runs(p, text)

            # Apply style colour + font + size + letter spacing to runs.
            size_pt = {1: s["docx_h1_size_pt"], 2: s["docx_h2_size_pt"], 3: s["docx_h3_size_pt"]}[level]
            color = {1: s["color_h1"], 2: s["color_h2"], 3: s["color_h3"]}[level]
            spacing_key = "h1_letter_spacing_pt" if level == 1 else "h2_letter_spacing_pt" if level == 2 else None
            spacing_pt = s.get(spacing_key, 0) if spacing_key else 0
            for run in p.runs:
                _apply_run_style(run, s["docx_heading_font"], size_pt, color)
                if spacing_pt:
                    _set_run_letter_spacing(run, spacing_pt)

            # Rules under headings
            if level == 1 and s.get("rule_under_h1"):
                _set_paragraph_bottom_border(p, s.get("rule_color", (200, 200, 200)), size=8)
            elif level == 2 and s.get("rule_under_h2"):
                rule_c = s.get("h2_rule_color", (200, 200, 200))
                _set_paragraph_bottom_border(p, rule_c, size=4)
            i += 1
            continue

        # Bullet list — consume consecutive bullet lines.
        if _BULLET_RE.match(line):
            while i < n:
                bm = _BULLET_RE.match(lines[i].rstrip())
                if not bm:
                    break
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, bm.group(1).strip())
                for run in p.runs:
                    if run.font.name != "Consolas":
                        run.font.name = s["docx_body_font"]
                    run.font.size = Pt(s["docx_body_size_pt"])
                    run.font.color.rgb = RGBColor(*s["color_body"])
                i += 1
            continue

        # Blockquote — render as italic paragraph.
        m = _BLOCKQUOTE_RE.match(line)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(1).strip())
            run.italic = True
            run.font.name = s["docx_body_font"]
            run.font.size = Pt(s["docx_body_size_pt"])
            run.font.color.rgb = RGBColor(*s["color_body"])
            i += 1
            continue

        # Plain paragraph — consume contiguous non-blank, non-special lines.
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if (
                _HEADING_RE.match(nxt)
                or _BULLET_RE.match(nxt)
                or _BLOCKQUOTE_RE.match(nxt)
                or _FENCE_RE.match(nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1
        text = " ".join(line_text.strip() for line_text in para_lines)
        p = doc.add_paragraph()
        _add_inline_runs(p, text)
        for run in p.runs:
            if run.font.name != "Consolas":
                run.font.name = s["docx_body_font"]
            run.font.size = Pt(s["docx_body_size_pt"])
            run.font.color.rgb = RGBColor(*s["color_body"])

    return doc


# ============================================================================
# PDF export (fpdf2). Walks the same parser as the docx exporter.
# Pure Python, no system dependencies — works on Windows, Streamlit Cloud, etc.
# Output is text-selectable so ATS scanners can read it.
# ============================================================================

_PDF_PARA_GAP = 2.0   # mm between paragraphs


def markdown_to_pdf_bytes(md: str, title: str | None = None, style: str | None = None) -> bytes:
    """Render Markdown to PDF bytes via Playwright (HTML/CSS -> Chromium -> PDF).

    Each style ships its own HTML/CSS template in core/html_templates.py. This
    gives proper designer-grade output (banners, two-column grids, accent colours,
    icons via Unicode) that fpdf2 could not match.
    """
    from . import html_templates  # noqa: PLC0415 — lazy to keep cold path fast
    style_key = (style or DEFAULT_STYLE) if (style or DEFAULT_STYLE) in html_templates._BUILDERS else "editorial_bold"
    blocks = _parse_md_blocks(md)
    html = html_templates.render_html(blocks, style_key)
    return _html_to_pdf_bytes(html, title=title)


def markdown_to_pdf(md: str, out_path: Path, title: str | None = None, style: str | None = None) -> Path:
    """Render Markdown to a .pdf file at out_path. Returns out_path."""
    data = markdown_to_pdf_bytes(md, title, style)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(data)
    return out_path


def _html_to_pdf_bytes(html: str, title: str | None = None) -> bytes:
    """Drive a headless Chromium via Playwright to render HTML -> PDF.

    Launches a fresh browser per call. For batch rendering (e.g. preview thumbnails
    of all four styles), prefer batch_html_to_pdf_bytes below to amortise startup.
    """
    from playwright.sync_api import sync_playwright  # noqa: PLC0415
    with sync_playwright() as p:
        browser = p.chromium.launch()
        try:
            page = browser.new_page()
            page.set_content(html, wait_until="networkidle", timeout=20000)
            return page.pdf(
                format="A4",
                print_background=True,
                prefer_css_page_size=True,
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
            )
        finally:
            browser.close()


def batch_html_to_pdf_bytes(htmls: dict[str, str]) -> dict[str, bytes]:
    """Render multiple HTML documents in a single Playwright session.

    Returns {key: pdf_bytes}. Used by preview thumbnails so a 4-style grid only
    spins up Chromium once instead of four times.
    """
    from playwright.sync_api import sync_playwright  # noqa: PLC0415
    results: dict[str, bytes] = {}
    with sync_playwright() as p:
        browser = p.chromium.launch()
        try:
            for key, html in htmls.items():
                page = browser.new_page()
                try:
                    page.set_content(html, wait_until="networkidle", timeout=20000)
                    results[key] = page.pdf(
                        format="A4",
                        print_background=True,
                        prefer_css_page_size=True,
                        margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
                    )
                finally:
                    page.close()
        finally:
            browser.close()
    return results


def _safe_text(text: str) -> str:
    """Strip characters outside cp1252 (Helvetica core font). Most resume text is fine."""
    return (text or "").encode("cp1252", errors="replace").decode("cp1252")


def _build_pdf(md: str, title: str | None, s: dict) -> FPDF:
    pdf = FPDF(unit="mm", format="A4")
    pdf.core_fonts_encoding = "cp1252"
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=18)
    if title:
        pdf.set_title(_safe_text(title))
    pdf.add_page()

    body_font = s["pdf_body_font"]
    heading_font = s["pdf_heading_font"]
    body_size = s["pdf_body_size"]
    line_h = s["pdf_line_h"]
    heading_sizes = {1: s["pdf_h1_size"], 2: s["pdf_h2_size"], 3: s["pdf_h3_size"]}
    color_h = {1: s["color_h1"], 2: s["color_h2"], 3: s["color_h3"]}

    # Full-page background fill (must be drawn first).
    page_fill = s.get("page_fill_color")
    if page_fill:
        pdf.set_fill_color(*page_fill)
        pdf.rect(0, 0, pdf.w, pdf.h, style="F")
        pdf.set_y(pdf.t_margin)

    # Top decorative horizontal rule (e.g. Grid Modern purple bar).
    top_rule = s.get("top_rule")
    if top_rule:
        pdf.set_draw_color(*top_rule["color"])
        pdf.set_line_width(top_rule.get("thickness_mm", 2.0))
        y = top_rule.get("y_offset_mm", 12)
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.set_line_width(0.2)
        if pdf.get_y() < y + 6:
            pdf.set_y(y + 6)

    # Draw banner at top of page before any content lands.
    banner = s.get("banner")
    if banner:
        pdf.set_fill_color(*banner["fill_color"])
        pdf.rect(0, 0, pdf.w, banner["height_mm"], style="F")
        pdf.set_y(banner.get("padding_top_mm", 10))
        pdf.set_x(pdf.l_margin)

    def _body():
        pdf.set_font(body_font, "", body_size)
        pdf.set_text_color(*s["color_body"])

    # If the layout is two-column, dispatch to that builder instead.
    if s.get("layout") == "two_column":
        _build_pdf_two_column(pdf, md, s, body_font, body_size, line_h,
                              heading_font, heading_sizes, color_h, _body)
        return pdf

    _body()

    lines = (md or "").splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip()

        if not line.strip():
            pdf.ln(_PDF_PARA_GAP)
            i += 1
            continue

        # Fenced code block — pass-through as a Courier paragraph.
        if _FENCE_RE.match(line):
            i += 1
            code_lines: list[str] = []
            while i < n and not _FENCE_RE.match(lines[i].rstrip()):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                pdf.set_font("Courier", "", max(body_size - 1, 9))
                for cl in code_lines:
                    pdf.write(line_h, _safe_text(cl))
                    pdf.ln(line_h)
                _body()
            if i < n:
                i += 1
            continue

        # Heading
        m = _HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 3)
            size = heading_sizes[level]
            pdf.ln(_PDF_PARA_GAP)
            pdf.set_font(heading_font, "B", size)
            pdf.set_text_color(*color_h[level])

            raw_text = m.group(2).strip()

            # Case transform for premium look (uppercase section headers).
            case_key = "h1_case" if level == 1 else "h2_case" if level == 2 else "h3_case"
            case = s.get(case_key, "normal")
            display = raw_text.upper() if case == "upper" else raw_text

            # Letter spacing on this heading.
            spacing_key = "h1_letter_spacing_pt" if level == 1 else "h2_letter_spacing_pt" if level == 2 else None
            spacing = s.get(spacing_key, 0) if spacing_key else 0
            if spacing:
                pdf.set_char_spacing(spacing * 0.35)

            # Render H1 with optional accent on first word OR centered alignment.
            align = s.get("h1_align", "left") if level == 1 else "left"
            if level == 1 and align == "center":
                # Centered single-color rendering. Skip accent-first-word in center mode.
                text_w = pdf.get_string_width(_safe_text(display))
                content_w = pdf.w - pdf.l_margin - pdf.r_margin
                offset = max(0, (content_w - text_w) / 2)
                pdf.set_x(pdf.l_margin + offset)
                pdf.write(size * 0.45, _safe_text(display))
            elif level == 1 and s.get("h1_accent_first_word") and " " in display:
                first_word, rest = display.split(" ", 1)
                accent = s.get("h1_accent_color", (232, 152, 39))
                pdf.set_text_color(*accent)
                pdf.write(size * 0.45, _safe_text(first_word + " "))
                pdf.set_text_color(*color_h[1])
                pdf.write(size * 0.45, _safe_text(rest))
            else:
                pdf.write(size * 0.45, _safe_text(display))
            if spacing:
                pdf.set_char_spacing(0)
            pdf.ln(size * 0.5)

            # Rules under headings
            rule_after = (level == 1 and s.get("rule_under_h1")) or (level == 2 and s.get("rule_under_h2"))
            if rule_after:
                rule_color = s.get("rule_color") if level == 1 else s.get("h2_rule_color", (200, 200, 200))
                if rule_color is None:
                    rule_color = (200, 200, 200)
                pdf.set_draw_color(*rule_color)
                thickness = s.get("rule_thickness", 0.4) if level == 1 else 0.2
                pdf.set_line_width(thickness)
                y = pdf.get_y() - 1
                pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
                pdf.set_line_width(0.2)  # restore
                pdf.ln(2.5 if level == 1 else 1.0)

            # If we just rendered H1 and the style has a banner, jump out of it.
            if level == 1 and s.get("banner"):
                banner_h = s["banner"]["height_mm"]
                if pdf.get_y() < banner_h + 4:
                    pdf.set_y(banner_h + 6)

            _body()
            i += 1
            continue

        # Bullet list (consume contiguous bullets)
        if _BULLET_RE.match(line):
            bullet = s["bullet"]
            while i < n:
                bm = _BULLET_RE.match(lines[i].rstrip())
                if not bm:
                    break
                _body()
                pdf.write(line_h, _safe_text(f"  {bullet}  "))
                _render_inline_pdf(pdf, bm.group(1).strip(), s)
                pdf.ln(line_h)
                i += 1
            continue

        # Blockquote → italic paragraph
        m = _BLOCKQUOTE_RE.match(line)
        if m:
            pdf.set_font(body_font, "I", body_size)
            pdf.set_text_color(*s["color_body"])
            pdf.write(line_h, _safe_text(m.group(1).strip()))
            pdf.ln(line_h)
            _body()
            i += 1
            continue

        # Plain paragraph — consume contiguous non-blank, non-special lines
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if (
                _HEADING_RE.match(nxt)
                or _BULLET_RE.match(nxt)
                or _BLOCKQUOTE_RE.match(nxt)
                or _FENCE_RE.match(nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1
        text = " ".join(s_line.strip() for s_line in para_lines)
        _render_inline_pdf(pdf, text, s)
        pdf.ln(line_h)

    return pdf


def _parse_md_blocks(md: str) -> list[tuple]:
    """Parse markdown into a list of (kind, payload) tuples for column routing.

    Kinds:
      ('h', level, text), ('p', text), ('bullets', [str, ...]), ('quote', text), ('code', text)
    """
    blocks: list[tuple] = []
    lines = (md or "").splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if _FENCE_RE.match(line):
            i += 1
            code_lines: list[str] = []
            while i < n and not _FENCE_RE.match(lines[i].rstrip()):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code_lines)))
            if i < n:
                i += 1
            continue
        m = _HEADING_RE.match(line)
        if m:
            blocks.append(("h", min(len(m.group(1)), 3), m.group(2).strip()))
            i += 1
            continue
        if _BULLET_RE.match(line):
            items: list[str] = []
            while i < n:
                bm = _BULLET_RE.match(lines[i].rstrip())
                if not bm:
                    break
                items.append(bm.group(1).strip())
                i += 1
            blocks.append(("bullets", items))
            continue
        m = _BLOCKQUOTE_RE.match(line)
        if m:
            blocks.append(("quote", m.group(1).strip()))
            i += 1
            continue
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if (_HEADING_RE.match(nxt) or _BULLET_RE.match(nxt)
                    or _BLOCKQUOTE_RE.match(nxt) or _FENCE_RE.match(nxt)):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append(("p", " ".join(s.strip() for s in para_lines)))
    return blocks


def _route_blocks_to_columns(blocks: list[tuple], sidebar_names: set[str]) -> tuple[list, list]:
    """Split parsed blocks into (main_blocks, sidebar_blocks) based on H2 section name."""
    main: list[tuple] = []
    sidebar: list[tuple] = []
    current = main  # Default before any H2 (e.g., H1 + intro paragraph)
    name_set = {n.lower() for n in sidebar_names}
    for blk in blocks:
        if blk[0] == "h" and blk[1] == 2:
            section_name = blk[2].strip().lower()
            current = sidebar if section_name in name_set else main
        current.append(blk)
    return main, sidebar


def _build_pdf_two_column(pdf: FPDF, md: str, s: dict, body_font: str, body_size: float,
                          line_h: float, heading_font: str, heading_sizes: dict,
                          color_h: dict, _body) -> None:
    """Render in two columns: main on left, sidebar on right. H1 spans full width."""
    blocks = _parse_md_blocks(md)
    sidebar_names = s.get("sidebar_sections", set())

    # H1 always renders full-width at top, before column split.
    if blocks and blocks[0][0] == "h" and blocks[0][1] == 1:
        h1_block = blocks.pop(0)
        _render_pdf_block(pdf, h1_block, s, body_font, body_size, line_h,
                          heading_font, heading_sizes, color_h, _body, full_width=True)
        # If the H1 just rendered inside a banner, jump cursor below the banner.
        banner = s.get("banner")
        if banner and pdf.get_y() < banner["height_mm"] + 4:
            pdf.set_y(banner["height_mm"] + 8)

    main, sidebar = _route_blocks_to_columns(blocks, sidebar_names)

    page_left = pdf.l_margin
    page_right = pdf.w - pdf.r_margin
    content_w = page_right - page_left
    sidebar_w = content_w * s.get("sidebar_width_ratio", 0.34)
    gutter = 6  # mm between columns
    main_w = content_w - sidebar_w - gutter
    main_x = page_left
    sidebar_x = page_left + main_w + gutter

    columns_start_y = pdf.get_y() + 2

    # Render main column
    pdf.set_left_margin(main_x)
    pdf.set_right_margin(pdf.w - (main_x + main_w))
    pdf.set_xy(main_x, columns_start_y)
    for blk in main:
        _render_pdf_block(pdf, blk, s, body_font, body_size, line_h,
                          heading_font, heading_sizes, color_h, _body)
    main_end_y = pdf.get_y()

    # Render sidebar column
    pdf.set_left_margin(sidebar_x)
    pdf.set_right_margin(pdf.w - (sidebar_x + sidebar_w))
    pdf.set_xy(sidebar_x, columns_start_y)
    for blk in sidebar:
        _render_pdf_block(pdf, blk, s, body_font, body_size, line_h,
                          heading_font, heading_sizes, color_h, _body)

    # Restore default margins for future pages
    pdf.set_left_margin(page_left)
    pdf.set_right_margin(pdf.w - page_right)


def _render_pdf_block(pdf: FPDF, blk: tuple, s: dict, body_font: str, body_size: float,
                     line_h: float, heading_font: str, heading_sizes: dict,
                     color_h: dict, _body, full_width: bool = False) -> None:
    """Render a single parsed block at the current pdf position."""
    kind = blk[0]
    if kind == "h":
        level = blk[1]
        raw_text = blk[2]
        size = heading_sizes[level]
        pdf.ln(_PDF_PARA_GAP)
        pdf.set_font(heading_font, "B", size)
        pdf.set_text_color(*color_h[level])

        case_key = "h1_case" if level == 1 else "h2_case" if level == 2 else "h3_case"
        display = raw_text.upper() if s.get(case_key, "normal") == "upper" else raw_text
        spacing_key = "h1_letter_spacing_pt" if level == 1 else "h2_letter_spacing_pt" if level == 2 else None
        spacing = s.get(spacing_key, 0) if spacing_key else 0
        if spacing:
            pdf.set_char_spacing(spacing * 0.35)

        align = s.get("h1_align", "left") if level == 1 else "left"
        if level == 1 and align == "center" and not full_width:
            # Centered within column. (When full_width=True for H1 we still center across page.)
            text_w = pdf.get_string_width(_safe_text(display))
            col_w = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.set_x(pdf.l_margin + max(0, (col_w - text_w) / 2))
            pdf.write(size * 0.45, _safe_text(display))
        elif level == 1 and align == "center" and full_width:
            text_w = pdf.get_string_width(_safe_text(display))
            col_w = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.set_x(pdf.l_margin + max(0, (col_w - text_w) / 2))
            pdf.write(size * 0.45, _safe_text(display))
        elif level == 1 and s.get("h1_accent_first_word") and " " in display:
            first_word, rest = display.split(" ", 1)
            accent = s.get("h1_accent_color", (232, 152, 39))
            pdf.set_text_color(*accent)
            pdf.write(size * 0.45, _safe_text(first_word + " "))
            pdf.set_text_color(*color_h[1])
            pdf.write(size * 0.45, _safe_text(rest))
        else:
            pdf.write(size * 0.45, _safe_text(display))
        if spacing:
            pdf.set_char_spacing(0)
        pdf.ln(size * 0.5)

        rule_after = (level == 1 and s.get("rule_under_h1")) or (level == 2 and s.get("rule_under_h2"))
        if rule_after:
            rc = s.get("rule_color") if level == 1 else s.get("h2_rule_color", (200, 200, 200))
            if rc is None:
                rc = (200, 200, 200)
            pdf.set_draw_color(*rc)
            thickness = s.get("rule_thickness", 0.4) if level == 1 else 0.2
            pdf.set_line_width(thickness)
            y = pdf.get_y() - 1
            # Line spans current column width (l_margin to w - r_margin since we adjusted margins)
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.set_line_width(0.2)
            pdf.ln(2.5 if level == 1 else 1.0)
        _body()
    elif kind == "p":
        _render_inline_pdf(pdf, blk[1], s)
        pdf.ln(line_h)
    elif kind == "bullets":
        bullet = s.get("bullet", "•")
        for item in blk[1]:
            _body()
            pdf.write(line_h, _safe_text(f"  {bullet}  "))
            _render_inline_pdf(pdf, item, s)
            pdf.ln(line_h)
    elif kind == "quote":
        pdf.set_font(body_font, "I", body_size)
        pdf.set_text_color(*s["color_body"])
        pdf.write(line_h, _safe_text(blk[1]))
        pdf.ln(line_h)
        _body()
    elif kind == "code":
        pdf.set_font("Courier", "", max(body_size - 1, 9))
        for cl in blk[1].splitlines():
            pdf.write(line_h, _safe_text(cl))
            pdf.ln(line_h)
        _body()


def _render_inline_pdf(pdf: FPDF, text: str, s: dict) -> None:
    """Walk **bold** / *italic* / `code` runs, switching fonts inline."""
    if not text:
        return
    body_font = s["pdf_body_font"]
    body_size = s["pdf_body_size"]
    line_h = s["pdf_line_h"]
    pdf.set_text_color(*s["color_body"])

    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            pdf.set_font(body_font, "B", body_size)
            pdf.write(line_h, _safe_text(part[2:-2]))
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            pdf.set_font(body_font, "I", body_size)
            pdf.write(line_h, _safe_text(part[1:-1]))
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            pdf.set_font("Courier", "", body_size)
            pdf.write(line_h, _safe_text(part[1:-1]))
        else:
            pdf.set_font(body_font, "", body_size)
            pdf.write(line_h, _safe_text(part))
    pdf.set_font(body_font, "", body_size)


def _add_inline_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph, splitting on **bold**, *italic*, `code`."""
    if not text:
        return
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)
